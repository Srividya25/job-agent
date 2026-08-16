"""Resume loading: PDF/DOCX -> text -> a comparable representation.

Parsed once and cached on disk. Nothing here is company-specific; the same
parsed resume is scored against every job.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path

from ..config import ROOT, ResumeRef, data_dir

_CACHE = "resume_cache.json"


@dataclass
class ParsedResume:
    label: str
    path: str
    text: str
    target_roles: list[str] = field(default_factory=list)

    @property
    def tokens(self) -> set[str]:
        return tokenize(self.text)


def _read_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def read_resume_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"resume not found: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _read_docx(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(errors="replace")
    raise ValueError(f"unsupported resume format: {suffix}")


_WORD = re.compile(r"[a-z0-9][a-z0-9+#.\-]{1,}")
_STOP = {
    "the", "and", "for", "with", "you", "our", "are", "will", "that", "this",
    "have", "from", "your", "who", "all", "can", "not", "but", "was", "has",
    "job", "work", "team", "role", "years", "year", "experience", "including",
    "using", "across", "within", "their", "them", "they", "more", "than",
    "into", "each", "such", "also", "other", "any", "how", "new", "may",
}


def tokenize(text: str) -> set[str]:
    return {
        w
        for w in _WORD.findall(text.lower())
        if w not in _STOP and len(w) > 2
    }


def load_resumes(refs: list[ResumeRef], refresh: bool = False) -> list[ParsedResume]:
    """Parse each configured resume, caching extracted text by mtime."""
    cache_path = data_dir() / _CACHE
    cache: dict = {}
    if cache_path.exists() and not refresh:
        try:
            cache = json.loads(cache_path.read_text())
        except json.JSONDecodeError:
            cache = {}

    out: list[ParsedResume] = []
    dirty = False

    for ref in refs:
        path = ref.resolve()
        if not path.exists():
            continue
        stamp = str(path.stat().st_mtime_ns)
        entry = cache.get(ref.label)

        if entry and entry.get("stamp") == stamp:
            text = entry["text"]
        else:
            text = read_resume_text(path)
            cache[ref.label] = {"stamp": stamp, "text": text}
            dirty = True

        out.append(
            ParsedResume(
                label=ref.label,
                path=str(path.relative_to(ROOT)),
                text=text,
                target_roles=ref.target_roles,
            )
        )

    if dirty:
        cache_path.write_text(json.dumps(cache))
    return out
